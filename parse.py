from datetime import datetime
import docx, aiohttp, asyncio, logging
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from aiogram.types import FSInputFile
from os import remove
from site_requests import sber_ast, etp_ets, zakupki_gov, roseltorg
from main import bot
import db


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink


async def coroutine_maker(sites, date, titles, user_id, session, time_slot, count):
    coroutines = []
    for site in sites:
        if site == "Сбер аст":
            count["Сбер аст"] = 0
            coroutines += [sber_ast.search(date, title, session, time_slot, count) for title in titles]
        elif site == "Фабрикант":
            count["Фабрикант"] = 0
            coroutines += [etp_ets.search(date, title, session, time_slot, count) for title in titles]
        elif site == "ЕИС закупки":
            count["ЕИС закупки"] = 0
            coroutines += [zakupki_gov.search(date, title, session, time_slot, count) for title in titles]
        elif site == "Росэлторг":
            count["Росэлторг"] = 0
            coroutines += [roseltorg.search(date, title, session, time_slot, count) for title in titles]
    return coroutines


async def parser(user_id, forcibly=False, state=None):
    date = datetime.now()
    user_data = db.users.find_one({"user_id": user_id})
    days = user_data["days"]
    if date.weekday() in days or forcibly:
        sites = user_data["sites"]
        key_words = user_data["keywords"]
        time_slot = user_data["filters"]["time_slot"]
        empty_doc = True
        doc = docx.Document()
        par = doc.add_paragraph(f'Выписка по актуальным тендерам на {date.strftime("%m.%d.%Y %H:%M")} за последние {time_slot} дня')
        await bot.send_message(user_id, "Начинаю проверку", disable_notification=True)
        async with aiohttp.ClientSession() as session:
            count = {}
            coroutines = await coroutine_maker(sites, date, key_words, user_id, session, time_slot, count)
            if forcibly:
                coroutines.append(progress_bar(count, len(key_words), user_id))
            results = await asyncio.gather(*coroutines)
        for site in results:
            for tender in site:
                empty_doc = False
                par = doc.add_paragraph(f'Номер: {tender["number"]}\nНаименование: {tender["name"]}\nЗаказчик: {tender["customer"]}\nЦена: {tender["price"]}\nОсталось времени на подачу: {tender["end_date"]}\n')
                add_hyperlink(par, "Подробнее", tender["more_info"])
        if empty_doc:
            await bot.send_message(chat_id=user_id, text="Нет новых тендеров")
        else:
            doc.save(f'{user_id}.docx')
            await bot.send_message(chat_id=user_id, text="Проверку закончил")
            await bot.send_document(user_id, document=FSInputFile(path=f'{user_id}.docx', filename=f'Актуальные тендеры {date.strftime("%m.%d.%Y - %H_%M")}.docx'))
            remove(f'{user_id}.docx')
        if state:
            await state.clear()


async def progress_bar(count, key_words_len, user_id):
    message_text = ""
    for elem in count:
        message_text += f'{elem}  ▱▱▱▱▱▱▱▱▱▱ 0%\n'
    message = await bot.send_message(user_id, message_text, disable_notification=True)
    old_text = message_text
    while any(value != key_words_len for value in count.values()):
        message_text = ""
        for elem in count:
            percent = round((count[elem] / key_words_len) * 100)
            progress_bar = "▰" * (round(percent, -1) // 10) + "▱" * (10 - round(percent, -1) // 10)
            message_text += f'{elem}  {progress_bar} {percent}%\n'
        if old_text != message_text:
            await bot.edit_message_text(message_text, user_id, message.message_id)
            old_text = message_text
        await asyncio.sleep(0.15)
    message_text = ""
    for elem in count:
        message_text += f'{elem}  ▰▰▰▰▰▰▰▰▰▰ 100%\n'
    if old_text != message_text:
        await bot.edit_message_text(message_text, user_id, message.message_id)
    return []